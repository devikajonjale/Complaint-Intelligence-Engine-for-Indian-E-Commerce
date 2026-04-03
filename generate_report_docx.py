"""Generate the academic-style project report as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    out = Path(__file__).resolve().parent / "Complaint_Intelligence_Engine_Project_Report.docx"
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Complaint Intelligence Engine for Indian E-Commerce")
    tr.bold = True
    tr.font.size = Pt(18)
    doc.add_paragraph()
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(
        "Unsupervised NLP Pipeline, Anomaly Detection, and Streamlit Dashboard"
    )
    sr.font.size = Pt(12)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Project Report (Word)").font.size = Pt(11)
    doc.add_page_break()

    # Abstract
    add_heading(doc, "Abstract", 1)
    doc.add_paragraph(
        "Indian e-commerce platforms generate large volumes of public reviews that mix English, "
        "Hindi, and Hinglish. Manually categorizing complaint themes and spotting urgent outliers is "
        "slow and inconsistent. This project implements a Complaint Intelligence Engine: an end-to-end "
        "unsupervised natural language processing (NLP) pipeline that ingests app-store reviews, "
        "cleans and enriches text, builds multilingual semantic embeddings, reduces dimensionality, "
        "clusters complaint archetypes, and detects anomalies and weekly complaint spikes. Results are "
        "exposed through a four-tab Streamlit dashboard for exploration, monitoring, and triage. "
        "The system is modular (numbered scripts), reproducible via a single orchestrator "
        "(run_pipeline.py), and designed for real-world constraints such as missing API credentials "
        "and timezone-aware timestamps in exported data."
    )

    # Objectives
    add_heading(doc, "Objectives", 1)
    add_bullets(
        doc,
        [
            "Ingest review data from multiple Indian e-commerce platforms (Google Play) and optionally Reddit.",
            "Preprocess noisy, multilingual text (including Hinglish) and engineer interpretable features.",
            "Represent reviews in a semantic vector space using multilingual Sentence-BERT.",
            "Discover complaint archetypes using unsupervised clustering (K-Means and HDBSCAN).",
            "Flag high-risk or unusual complaints using Isolation Forest and One-Class SVM, with Tier-1 critical labeling.",
            "Track weekly complaint volume and highlight spikes using rolling z-score logic.",
            "Deliver an interactive Streamlit demo for non-technical stakeholders.",
        ],
    )

    # Methodology
    add_heading(doc, "Methodology", 1)
    add_heading(doc, "Data collection", 2)
    doc.add_paragraph(
        "Reviews are collected using the Google Play scraper for configured apps (Myntra, Meesho, Nykaa, "
        "Flipkart, Amazon India) across English and Hindi listings where available. Reddit posts may be "
        "merged when PRAW credentials are supplied via a .env file; otherwise the pipeline continues "
        "with Play Store data only."
    )
    add_heading(doc, "Preprocessing and features", 2)
    doc.add_paragraph(
        "Text is normalized, deduplicated, and filtered. Language detection identifies mixed-language "
        "content; Hinglish is flagged. Star ratings are bucketed (e.g., complaint vs positive). "
        "A TF-IDF matrix supports lexical analysis alongside dense embeddings."
    )
    add_heading(doc, "Embeddings and dimensionality reduction", 2)
    doc.add_paragraph(
        "Sentence embeddings are produced with the sentence-transformers model "
        "paraphrase-multilingual-MiniLM-L12-v2 (384 dimensions). PCA reduces variance to a manageable "
        "subspace; UMAP yields a 10D manifold-friendly representation; t-SNE provides a 2D layout for "
        "visualization in the dashboard."
    )
    add_heading(doc, "Clustering and evaluation", 2)
    doc.add_paragraph(
        "K-Means clusters are chosen with silhouette-based k selection. HDBSCAN offers density-based "
        "grouping. Cluster profiles summarize dominant themes. Agreement between K-Means and HDBSCAN "
        "can be summarized via the Adjusted Rand Index (ARI) when both are available."
    )
    add_heading(doc, "Anomaly detection and spikes", 2)
    doc.add_paragraph(
        "Isolation Forest and One-Class SVM score unusualness relative to complaint and positive "
        "regions. Tier-1 critical alerts highlight the most severe combined anomalies. Weekly complaint "
        "counts by platform and cluster feed a spike detector based on rolling z-scores."
    )

    # Architecture / Workflow
    add_heading(doc, "Project Architecture and Workflow", 1)
    doc.add_paragraph(
        "The pipeline is a linear sequence of scripts, each reading/writing under data/ and models/, "
        "orchestrated by run_pipeline.py:"
    )
    add_bullets(
        doc,
        [
            "01_ingest.py — fetch and merge review sources; write raw_reviews.csv.",
            "02_preprocess.py — clean text, features, TF-IDF; write cleaned_reviews.csv and tfidf artifacts.",
            "03_embed.py — Sentence-BERT embeddings; write embeddings.npy.",
            "04_reduce.py — PCA, UMAP, t-SNE; persist reduced arrays and reducers.",
            "05_cluster.py — K-Means, HDBSCAN, profiles; clustered outputs.",
            "06_anomaly.py — anomaly scores, Tier-1 flags, spike_report.csv, final_reviews.csv.",
            "07_visualize.py — export static figures under figures/.",
            "app.py — Streamlit UI consuming final_reviews.csv, spike_report.csv, and tsne_2d.npy.",
        ],
    )
    doc.add_paragraph(
        "A high-level flow: Data → Preprocess → Embed → Reduce → Cluster → Anomaly & spikes → "
        "Visualize → Dashboard."
    )

    # Tech stack
    add_heading(doc, "Technology Stack", 1)
    add_bullets(
        doc,
        [
            "Language: Python 3.x",
            "Data: pandas, numpy",
            "ML: scikit-learn (PCA, K-Means, Isolation Forest, One-Class SVM, TF-IDF), "
            "umap-learn, hdbscan, scipy",
            "Deep learning / NLP: PyTorch, sentence-transformers, transformers, langdetect",
            "Visualization: plotly, matplotlib, seaborn",
            "Application: Streamlit",
            "Utilities: python-dotenv, joblib, google-play-scraper, praw (optional)",
        ],
    )

    # Limitations and mitigations
    add_heading(doc, "Limitations and Mitigations", 1)
    add_heading(doc, "Data and API constraints", 2)
    add_bullets(
        doc,
        [
            "Limitation: Reddit API may return 401 without client credentials. "
            "Mitigation: Optional .env (REDDIT_CLIENT_ID, REDDIT_CLIENT_SECRET); pipeline proceeds with Play Store data.",
            "Limitation: Play Store scrape limits and rate behavior. "
            "Mitigation: Batch pulls and merge; quality checks in preprocessing.",
        ],
    )
    add_heading(doc, "Modeling constraints", 2)
    add_bullets(
        doc,
        [
            "Limitation: One-Class SVM may lack sufficient “normal” positive rows in some slices. "
            "Mitigation: Fallback defaults and logging in the anomaly step.",
            "Limitation: Language detection can misclassify short or noisy text. "
            "Mitigation: Length filters, Hinglish heuristics, and star-bucket context.",
        ],
    )
    add_heading(doc, "Engineering and deployment", 2)
    add_bullets(
        doc,
        [
            "Limitation: Large dependencies (torch, transformers) and Windows file locks during pip install. "
            "Mitigation: Virtual environment, retry installs, optional Developer Mode for Hugging Face cache symlinks.",
            "Limitation: Streamlit date filter compared UTC-aware timestamps to naive dates (TypeError). "
            "Mitigation: Compare using UTC-aligned pd.Timestamp bounds when the at column is timezone-aware.",
        ],
    )

    # Impact / Applications
    add_heading(doc, "Impact and Applications", 1)
    add_bullets(
        doc,
        [
            "Customer experience teams: prioritize recurring complaint themes by cluster and platform.",
            "Trust and safety / ops: review Tier-1 anomaly-flagged reviews first.",
            "Product and category managers: correlate spikes with releases or campaigns using weekly spike charts.",
            "Research and education: template for unsupervised NLP on Indian e-commerce text with a clear dashboard.",
        ],
    )

    # Conclusion
    add_heading(doc, "Conclusion", 1)
    doc.add_paragraph(
        "The Complaint Intelligence Engine delivers a practical, end-to-end workflow from live review "
        "ingestion to interactive visualization. Multilingual embeddings and clustering surface "
        "complaint archetypes without manual labeling; anomaly detection and spike tracking add "
        "operational signals. The Streamlit application makes the outputs accessible to stakeholders "
        "who are not ML specialists. The pipeline is structured for extension (new sources, supervised "
        "labels, or API integration) and for reproducible runs via run_pipeline.py."
    )

    # Future scope
    add_heading(doc, "Future Scope", 1)
    add_bullets(
        doc,
        [
            "Integrate authenticated Reddit and additional social sources with robust rate limiting.",
            "Add automated tests and CI for each pipeline stage.",
            "Ship a frozen sample dataset for offline demos and grading.",
            "Fine-tune or distill a domain-specific embedding model for Indian e-commerce.",
            "Add explainability (e.g., representative phrases per cluster) and exportable PDF reports.",
            "Deploy on Streamlit Cloud or container with secrets management and scheduled pipeline runs.",
        ],
    )

    # Streamlit Demo
    add_heading(doc, "Streamlit Demo", 1)
    doc.add_paragraph(
        "The dashboard (app.py) loads processed outputs from the data/ folder after "
        "python run_pipeline.py has been executed at least once."
    )
    add_heading(doc, "How to run", 2)
    add_bullets(
        doc,
        [
            "Install dependencies: pip install -r requirements.txt",
            "Generate data: python run_pipeline.py",
            "Launch UI: streamlit run app.py",
            "Open the local URL shown in the terminal (typically http://localhost:8501).",
        ],
    )
    add_heading(doc, "Sidebar controls", 2)
    add_bullets(
        doc,
        [
            "Platform and Star Bucket multi-select filters.",
            "Date range filter (UTC-safe when timestamps are timezone-aware).",
            "Refresh Full Pipeline (re-runs run_pipeline.py; may take several minutes).",
        ],
    )
    add_heading(doc, "Tabs", 2)
    add_bullets(
        doc,
        [
            "Live Pulse — KPI metrics and stacked bar chart of complaint volume by platform and cluster.",
            "Complaint Landscape — interactive t-SNE scatter plot (cluster-colored, hover shows review text).",
            "Spike Tracker — weekly complaint line chart with red markers for detected spikes.",
            "Critical Alerts — table of Tier-1 complaints with isolation scores and truncated content.",
        ],
    )
    doc.add_paragraph(
        "Suggested demo path: filter to Star Bucket “complaint” and one platform (e.g., Myntra), "
        "review Live Pulse, explore Complaint Landscape tooltips, then open Critical Alerts for "
        "high-severity rows."
    )

    # References / outputs (optional short section)
    add_heading(doc, "Key Output Artifacts", 1)
    doc.add_paragraph(
        "Typical generated files include data/final_reviews.csv, data/spike_report.csv, "
        "models/*.pkl, figures/*.png, and data/embeddings.npy. The dashboard reads "
        "final_reviews.csv, spike_report.csv, and tsne_2d.npy for coordinates."
    )

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
